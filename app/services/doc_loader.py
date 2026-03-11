import os
# 关键：彻底移除 unstructured 的导入，改用 python-docx
try:
    from docx import Document
except ImportError:
    print("正在安装必要库: python-docx...")
    os.system("pip install python-docx")
    from docx import Document

class DocumentLoader:
    @staticmethod
    def load_document(file_path: str) -> str:
        """
        稳健版解析：仅使用 python-docx，避开 NumPy/onnxruntime 冲突
        """
        if not os.path.exists(file_path):
            return f"错误：找不到文件 {file_path}"

        ext = os.path.splitext(file_path)[-1].lower()
        print(f"[Loader] 正在读取: {os.path.basename(file_path)}")

        try:
            if ext == ".docx":
                doc = Document(file_path)
                full_text = []
                
                # 1. 提取段落
                for para in doc.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text.strip())
                
                # 2. 提取表格并转化为简单的 Markdown 样式
                for table in doc.tables:
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        full_text.append(" | ".join(row_data))
                
                return "\n\n".join(full_text)
            
            elif ext == ".pdf":
                # 如果以后非要用 PDF，建议环境修复后再加回来
                return "当前环境 NumPy 冲突，暂不支持 PDF，请优先使用 docx"
                
            else:
                return f"不支持的格式: {ext}"

        except Exception as e:
            # 这里会捕获 'File is not a zip file' 等错误
            return f"解析失败: {str(e)}"